"""产品手册解析模块：从 Word(.docx) 文件中提取结构化产品信息。

对应需求加分项：「支持上传产品手册（PDF/Word），自动解析并构建知识库」。

整体思路：
1. 用 python-docx 读取 .docx 的段落（含标题层级）和表格，拼出纯文本
2. 把纯文本交给"提取器"转成结构化产品字段：
   - LLM 启用时：让大模型从手册文本中抽取 JSON（最准确）
   - Mock 模式：用关键词/标题层级的启发式规则抽取（无需联网）

为什么解析后不直接入库？
→ 手册写法千差万别，自动抽取不可能 100% 准确。
   返回"草稿"让用户在前端核对、补全后再保存，比直接入库更可靠。

PDF 解析同理（可用 pdfplumber/PyMuPDF），本模块先实现 .docx，
后续如需支持 PDF，只需新增一个 extract_text_from_pdf 即可，
parse_product 的下游逻辑可完全复用。
"""
import io
import json
import os
import re
from typing import Any

from docx import Document

from app.agents.llm_provider import LLMProvider, get_provider


# ===== 第一层：从 .docx 读取原始文本 =====

def extract_blocks_from_docx(file_bytes: bytes) -> list[dict]:
    """读取 .docx 字节流，返回结构化块列表。

    每个块是字典：
    - {"type": "heading", "level": 1, "text": "..."}   标题（level 越小越靠前）
    - {"type": "paragraph", "text": "..."}             普通段落
    - {"type": "table", "text": "单元格 | 单元格 ..."} 表格（拍平成文本行）

    保留块的类型和顺序，便于启发式提取器判断"标题下的段落就是功能说明"。
    """
    # BytesIO 把字节流包装成"文件对象"，python-docx 才能读取
    doc = Document(io.BytesIO(file_bytes))
    blocks: list[dict] = []

    # paragraphs 是文档正文段落（不含表格内文字）
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        # style.name 形如 "Heading 1" / "标题 2" / "Normal"
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading") or style.startswith("标题"):
            # 从样式名提取层级数字，取不到默认 1
            m = re.search(r"(\d+)", style)
            level = int(m.group(1)) if m else 1
            blocks.append({"type": "heading", "level": level, "text": text})
        else:
            blocks.append({"type": "paragraph", "text": text})

    # tables 是文档中的表格，每个 table 拍平成多行文本
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                blocks.append({"type": "table", "text": line})

    return blocks


def blocks_to_text(blocks: list[dict]) -> str:
    """把块列表拼成一段纯文本（用于喂给 LLM 或统计字符数）。"""
    parts = []
    for b in blocks:
        if b["type"] == "heading":
            parts.append(f"## {b['text']}")
        else:
            parts.append(b["text"])
    return "\n".join(parts)


# ===== 第一层（PDF）：从 PDF 读取原始文本 =====

def extract_blocks_from_pdf(file_bytes: bytes) -> list[dict]:
    """读取 PDF 字节流，返回结构化块列表（与 extract_blocks_from_docx 同构）。

    使用 pdfplumber 逐页提取文本，按行拆分为块。
    PDF 没有像 Word 那样的标题样式信息，因此用启发式判断：
    - 短行（< 30 字）且不以标点结尾 -> 当作标题（heading）
    - 其余行 -> 普通段落（paragraph）

    这样 _extract_heuristic 仍能按"标题+紧邻段落"组织功能项。
    """
    import pdfplumber

    blocks: list[dict] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # 启发式：短行且不以标点结尾 -> 标题
                if len(line) < 30 and line[-1] not in "。.，,；;：:！!？?":
                    blocks.append({"type": "heading", "level": 2, "text": line})
                else:
                    blocks.append({"type": "paragraph", "text": line})
    return blocks


# ===== 第二层：结构化提取 =====

# 发给 LLM 的提示词：要求严格输出 JSON，字段对齐 ProductBase
_EXTRACT_PROMPT = """你是产品信息提取助手。从下面这份产品手册文本中，提取结构化产品信息。

手册文本：
{doc_text}

请输出 JSON，包含以下字段（缺失的字段给空值，不要编造）：
{{
  "name": "产品名称",
  "category": "产品类别",
  "tagline": "一句话标语（不超过 20 字）",
  "description": "产品详细描述（50-200 字）",
  "features": [{{"name": "功能名", "description": "功能说明"}}],
  "tech_params": [{{"name": "参数名", "value": "参数值"}}],
  "target_customers": ["目标客户行业"],
  "pricing": "定价信息",
  "competitors": ["竞品名称"],
  "selling_points": ["核心卖点关键词"]
}}

只输出 JSON，不要其他内容。"""


def _parse_blocks(
    blocks: list[dict], provider: LLMProvider
) -> dict[str, Any]:
    """通用解析：从块列表提取结构化产品信息（docx / pdf 共用）。

    返回字典结构：
    {
        "product": { ...ProductBase 字段... },
        "char_count": 文档字符数,
        "extractor": "llm" / "heuristic" / "none",
        "note": "解析说明"
    }
    """
    text = blocks_to_text(blocks)
    char_count = len(text)

    # 空文档直接返回空草稿
    if char_count == 0:
        return {
            "product": _empty_product(),
            "char_count": 0,
            "extractor": "none",
            "note": "文档内容为空，请手动录入。",
        }

    # 真实 LLM：让大模型抽取
    if not provider.is_mock:
        product, note = _extract_with_llm(provider, text)
        return {
            "product": product,
            "char_count": char_count,
            "extractor": "llm",
            "note": note,
        }

    # Mock 模式：启发式规则抽取
    product, note = _extract_heuristic(blocks, text)
    return {
        "product": product,
        "char_count": char_count,
        "extractor": "heuristic",
        "note": note,
    }


def parse_product_from_manual(
    file_bytes: bytes, filename: str = "", provider: LLMProvider | None = None
) -> dict[str, Any]:
    """解析产品手册（支持 PDF / Word），返回结构化产品草稿 + 元信息。

    根据文件扩展名自动选择解析器：
    - .pdf  -> extract_blocks_from_pdf（pdfplumber）
    - .docx -> extract_blocks_from_docx（python-docx）
    - 其余  -> 按 .docx 处理

    下游提取逻辑（LLM / 启发式）完全复用，无需为每种格式重复实现。
    """
    if provider is None:
        provider = get_provider()

    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        blocks = extract_blocks_from_pdf(file_bytes)
    else:
        blocks = extract_blocks_from_docx(file_bytes)
    return _parse_blocks(blocks, provider)


def parse_product_from_docx(
    file_bytes: bytes, provider: LLMProvider | None = None
) -> dict[str, Any]:
    """解析 .docx（向后兼容入口，等价于 parse_product_from_manual 不传 filename）。"""
    if provider is None:
        provider = get_provider()
    blocks = extract_blocks_from_docx(file_bytes)
    return _parse_blocks(blocks, provider)


def _extract_with_llm(provider: LLMProvider, text: str) -> tuple[dict, str]:
    """用 LLM 从手册文本中提取结构化产品信息。

    返回 (product_dict, note)
    """
    # 文本过长时截断，避免超出 LLM 上下文（保留开头，开头通常含产品名/定位）
    truncated = text[:4000]
    prompt = _EXTRACT_PROMPT.format(doc_text=truncated)
    raw = provider.chat(
        system_prompt="你是产品信息提取助手。只输出 JSON，不要其他内容。",
        user_prompt=prompt,
        temperature=0.1,    # 极低温度 → 抽取结果更稳定、不发散
    )
    try:
        obj = _parse_json(raw)
        product = _normalize_product(obj)
        return product, "已由 LLM 从手册中提取，请核对后保存。"
    except (json.JSONDecodeError, ValueError):
        # LLM 返回异常 → 回退到启发式
        blocks = _text_to_blocks(truncated)
        product, note = _extract_heuristic(blocks, truncated)
        return product, "LLM 解析失败，已回退到规则提取。" + note


def _extract_heuristic(blocks: list[dict], text: str) -> tuple[dict, str]:
    """启发式抽取：按"标题 → 章节正文"的结构化方式解析。

    核心思路：
    1. 把 blocks 按 heading 切分成多个 section（标题 + 其后段落列表）
    2. 根据每个 section 的标题关键词，映射到对应产品字段：
       - 名称/产品名 → product.name
       - 分类/类别 → category
       - 描述/简介 → description
       - 功能/特性 → features（章节正文每行 = 一个功能）
       - 目标客户/适用客户 → target_customers
       - 价格/定价 → pricing
       - 卖点/优势 → selling_points
       - 竞品/对比 → competitors
       - 技术参数/规格 → tech_params
    3. 无 heading 结构时回退到关键词逐行扫描（仅扫描段落，不含标题本身）
    """
    product = _empty_product()
    found: list[str] = []

    # ===== 章节标题正则 =====
    _RE_NAME = re.compile(r"产品名称|产品名|^名称$")
    _RE_CATEGORY = re.compile(r"产品分类|^分类$|类别|类目|所属分类")
    _RE_DESC = re.compile(r"产品描述|^描述$|产品简介|^简介$|产品介绍|产品概述|^概述$")
    _RE_FEATURE = re.compile(r"功能特性|核心功能|^功能$|功能列表|主要功能|功能介绍|产品功能|特性")
    _RE_TARGET = re.compile(r"目标客户|适用客户|客户群体|目标人群|适用行业|目标对象|^客户$")
    _RE_PRICING = re.compile(r"^价格$|价格信息|定价|费用|报价|收费|套餐")
    _RE_SELLING = re.compile(r"^卖点$|产品卖点|核心优势|^优势$|亮点|特点|价值主张")
    _RE_COMPETITOR = re.compile(r"竞品|竞争分析|竞品对比|^对比$|竞对")
    _RE_TECH = re.compile(r"技术参数|技术规格|^参数$|^规格$|技术指标")

    # ===== 1. 切分章节 =====
    sections: list[tuple[str, list[str]]] = []
    cur_heading = ""
    cur_lines: list[str] = []
    for b in blocks:
        if b["type"] == "heading":
            if cur_heading or cur_lines:
                sections.append((cur_heading, cur_lines))
            cur_heading = b["text"].strip()
            cur_lines = []
        elif b["type"] == "paragraph":
            cur_lines.append(b["text"].strip())
        elif b["type"] == "table":
            cur_lines.append(b["text"].strip())
    if cur_heading or cur_lines:
        sections.append((cur_heading, cur_lines))

    # ===== 2. 从第一个章节确定产品名 =====
    if sections:
        first_h, first_lines = sections[0]
        if _RE_NAME.search(first_h) and first_lines:
            # 标题是"产品名称"，实际名称在后面的段落
            product["name"] = first_lines[0][:50]
            found.append("名称")
            sections = sections[1:]   # 已消费，后续不再处理
        elif first_h and not _RE_NAME.search(first_h):
            # 第一个标题本身就是产品名
            product["name"] = first_h[:50]
            found.append("名称")
        elif first_lines:
            product["name"] = first_lines[0][:20]
            found.append("名称")

    # ===== 3. 按章节标题映射到字段 =====
    description_parts: list[str] = []
    features: list[dict] = []
    target_customers: list[str] = []
    selling_points: list[str] = []
    competitors: list[str] = []
    tech_params: list[dict] = []
    pricing_lines: list[str] = []
    category_vals: list[str] = []

    for heading, lines in sections:
        section_text = " ".join(l for l in lines if l).strip()

        # 产品分类
        if _RE_CATEGORY.search(heading):
            content = _strip_label_prefix(section_text)
            if content:
                for c in re.split(r"[，,、；;\n]", content):
                    c = c.strip()
                    if c and c not in category_vals:
                        category_vals.append(c)
            continue

        # 产品描述
        if _RE_DESC.search(heading):
            if section_text:
                description_parts.append(section_text)
            continue

        # 功能特性：每行 = 一个功能
        if _RE_FEATURE.search(heading):
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                name, desc = _split_feature_line(line)
                if name:
                    features.append({"name": name[:30], "description": desc[:120]})
            continue

        # 目标客户
        if _RE_TARGET.search(heading):
            content = _strip_label_prefix(section_text)
            if content:
                for tc in re.split(r"[，,、；;\n]", content):
                    tc = tc.strip()
                    if 2 <= len(tc) <= 30 and tc not in target_customers:
                        target_customers.append(tc)
            continue

        # 价格
        if _RE_PRICING.search(heading):
            if section_text:
                pricing_lines.append(section_text[:200])
            continue

        # 卖点
        if _RE_SELLING.search(heading):
            content = _strip_label_prefix(section_text)
            if content:
                for sp in re.split(r"[，,、；;\n]", content):
                    sp = sp.strip()
                    if 2 <= len(sp) <= 30 and sp not in selling_points:
                        selling_points.append(sp)
            continue

        # 竞品
        if _RE_COMPETITOR.search(heading):
            content = _strip_label_prefix(section_text)
            if content:
                for c in re.split(r"[，,、和与；;\n]", content):
                    c = c.strip()
                    if 2 <= len(c) <= 20 and c not in competitors:
                        competitors.append(c)
            continue

        # 技术参数
        if _RE_TECH.search(heading):
            for line in lines:
                line = line.strip()
                if ("：" in line or ":" in line) and len(line) < 80:
                    sep = "：" if "：" in line else ":"
                    k, _, v = line.partition(sep)
                    if k.strip() and v.strip():
                        tech_params.append({"name": k.strip()[:20], "value": v.strip()[:40]})
            continue

    # ===== 4. 回退：无章节命中时，对段落逐行关键词扫描 =====
    para_lines = [b["text"].strip() for b in blocks if b["type"] == "paragraph"]
    if not description_parts and para_lines:
        description_parts.append(" ".join(para_lines[:3])[:300])
    for line in para_lines:
        # 定价回退
        if not pricing_lines and re.search(r"[¥￥]|\d+元|定价|价格|报价|/年|/月", line) and len(line) < 200:
            pricing_lines.append(line[:200])
        # 卖点回退
        if not selling_points and re.search(r"优势|卖点|亮点|特点", line) and len(line) < 100:
            content = _strip_label_prefix(line)
            for sp in re.split(r"[，,、；;]", content):
                sp = sp.strip()
                if 2 <= len(sp) <= 30 and sp not in selling_points:
                    selling_points.append(sp)
        # 目标客户回退
        if not target_customers and re.search(r"目标客户|适用客户|客户群体|适用行业", line) and len(line) < 100:
            content = _strip_label_prefix(line)
            for tc in re.split(r"[，,、；;\n]", content):
                tc = tc.strip()
                if 2 <= len(tc) <= 30 and tc not in target_customers:
                    target_customers.append(tc)
        # 竞品回退
        if not competitors and re.search(r"竞品|对比|竞对", line):
            content = _strip_label_prefix(line)
            for c in re.split(r"[，,、和与]", content):
                c = c.strip()
                if 2 <= len(c) <= 20 and c not in competitors:
                    competitors.append(c)

    # ===== 5. 赋值到 product =====
    if description_parts:
        product["description"] = " ".join(description_parts)[:500]
        found.append("描述")
    if category_vals:
        product["category"] = category_vals[:5]
        found.append("分类")
    if features:
        product["features"] = features[:8]
        found.append("功能")
    if tech_params:
        product["tech_params"] = tech_params[:10]
        found.append("技术参数")
    if pricing_lines:
        product["pricing"] = "；".join(pricing_lines)[:300]
        found.append("定价")
    if target_customers:
        product["target_customers"] = target_customers[:8]
        found.append("目标客户")
    if selling_points:
        product["selling_points"] = selling_points[:6]
        found.append("卖点")
    if competitors:
        product["competitors"] = competitors[:6]
        found.append("竞品")

    note = f"规则提取完成，识别到：{ '、'.join(found) if found else '未识别到关键字段' }。请核对补全后保存。"
    return product, note


def _strip_label_prefix(text: str) -> str:
    """去掉"标签："或"标签:"前缀，返回纯内容。"""
    for sep in ("：", ":"):
        if sep in text:
            parts = text.split(sep, 1)
            # 仅当前缀较短（像标签而非正文中的冒号）时才去掉
            if len(parts[0]) <= 10:
                return parts[1].strip()
    return text.strip()


def _split_feature_line(line: str) -> tuple[str, str]:
    """把一行功能描述拆成 (功能名, 描述)。

    支持格式：
    - "功能名：描述"
    - "功能名: 描述"
    - "功能名 - 描述"
    - "- 功能名：描述"（列表项）
    - "功能名"（无描述）
    """
    line = re.sub(r"^[\s•·\-\*]+", "", line).strip()    # 去掉列表符号
    for sep in ("：", ":", " - ", " – ", " — "):
        if sep in line:
            name, _, desc = line.partition(sep)
            name = name.strip()
            desc = desc.strip()
            if name and len(name) <= 30:
                return name, desc
    return line[:30], ""


# ===== 工具函数 =====

def _empty_product() -> dict:
    """返回空的产品草稿（字段结构与 ProductBase 对齐）。"""
    return {
        "name": "",
        "category": [],
        "tagline": "",
        "description": "",
        "features": [],
        "tech_params": [],
        "target_customers": [],
        "pricing": "",
        "competitors": [],
        "selling_points": [],
    }


def _normalize_product(obj: dict) -> dict:
    """把 LLM 返回的 JSON 规范化为产品草稿。

    主要是保证各字段类型正确（LLM 偶尔会把数组返回成字符串），
    缺失字段补默认空值。
    """
    base = _empty_product()
    base.update({k: v for k, v in obj.items() if k in base})
    # category：LLM 常返回字符串，需转为列表以匹配 ProductBase.category: list[str]
    cat = base.get("category")
    if isinstance(cat, str) and cat.strip():
        base["category"] = [cat.strip()]
    elif not isinstance(cat, list):
        base["category"] = []
    # 确保列表字段确实是列表
    for key in ("features", "tech_params", "target_customers", "competitors", "selling_points"):
        if not isinstance(base.get(key), list):
            base[key] = []
    # features / tech_params 元素补全键
    base["features"] = [
        {"name": f.get("name", ""), "description": f.get("description", "")}
        for f in base["features"] if isinstance(f, dict)
    ]
    base["tech_params"] = [
        {"name": t.get("name", ""), "value": t.get("value", "")}
        for t in base["tech_params"] if isinstance(t, dict)
    ]
    return base


def _parse_json(raw: str) -> dict:
    """解析 LLM 返回的 JSON，兼容 ```json 代码块包裹。"""
    text = raw.strip()
    if text.startswith("```"):
        # 去掉首尾的 ``` 围栏
        text = text.split("```", 2)
        # text[1] 是围栏内的内容，可能以 "json" 开头
        inner = text[1] if len(text) > 1 else ""
        if inner.startswith("json"):
            inner = inner[4:]
        text = inner.strip()
    return json.loads(text)


def _text_to_blocks(text: str) -> list[dict]:
    """纯文本回退成块列表（LLM 失败时用）。

    把以 ## 开头的行当作标题，其余当作段落。
    """
    blocks = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            blocks.append({"type": "heading", "level": 2, "text": line[3:].strip()})
        else:
            blocks.append({"type": "paragraph", "text": line})
    return blocks
