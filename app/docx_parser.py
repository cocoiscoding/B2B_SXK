"""产品手册解析模块：从 Word(.docx) 文件中提取结构化产品信息。

对应需求加分项：「支持上传产品手册（PDF/Word），自动解析并构建知识库」。

整体思路：
1. 用 python-docx 读取 .docx 的段落（含标题层级）和表格，拼出纯文本
2. 把纯文本交给"提取器"转成结构化产品字段：
   - LLM 启用时：让大模型从手册文本中抽取 JSON（最准确）
   - Mock 模式：用关键词/标题层级的启发式规则抽取（无需联网）

为什么解析后不直接入库？
→ 手册写法千差万别，自动抽取不可能 100% 准确。
   返回"草稿"让用户在前端核对、补全后再保存，比直接入库更可靠。

PDF 解析同理（可用 pdfplumber/PyMuPDF），本模块先实现 .docx，
后续如需支持 PDF，只需新增一个 extract_text_from_pdf 即可，
parse_product 的下游逻辑可完全复用。
"""
import io
import json
import re
from typing import Any

from docx import Document

from app.agents.llm_provider import LLMProvider, get_provider


# ===== 第一层：从 .docx 读取原始文本 =====

def extract_blocks_from_docx(file_bytes: bytes) -> list[dict]:
    """读取 .docx 字节流，返回结构化块列表。

    每个块是字典：
    - {"type": "heading", "level": 1, "text": "..."}   标题（level 越小越靠前）
    - {"type": "paragraph", "text": "..."}             普通段落
    - {"type": "table", "text": "单元格 | 单元格 ..."} 表格（拍平成文本行）

    保留块的类型和顺序，便于启发式提取器判断"标题下的段落就是功能说明"。
    """
    # BytesIO 把字节流包装成"文件对象"，python-docx 才能读取
    doc = Document(io.BytesIO(file_bytes))
    blocks: list[dict] = []

    # paragraphs 是文档正文段落（不含表格内文字）
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        # style.name 形如 "Heading 1" / "标题 2" / "Normal"
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading") or style.startswith("标题"):
            # 从样式名提取层级数字，取不到默认 1
            m = re.search(r"(\d+)", style)
            level = int(m.group(1)) if m else 1
            blocks.append({"type": "heading", "level": level, "text": text})
        else:
            blocks.append({"type": "paragraph", "text": text})

    # tables 是文档中的表格，每个 table 拍平成多行文本
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                blocks.append({"type": "table", "text": line})

    return blocks


def blocks_to_text(blocks: list[dict]) -> str:
    """把块列表拼成一段纯文本（用于喂给 LLM 或统计字符数）。"""
    parts = []
    for b in blocks:
        if b["type"] == "heading":
            parts.append(f"## {b['text']}")
        else:
            parts.append(b["text"])
    return "\n".join(parts)


# ===== 第二层：结构化提取 =====

# 发给 LLM 的提示词：要求严格输出 JSON，字段对齐 ProductBase
_EXTRACT_PROMPT = """你是产品信息提取助手。从下面这份产品手册文本中，提取结构化产品信息。

手册文本：
{doc_text}

请输出 JSON，包含以下字段（缺失的字段给空值，不要编造）：
{{
  "name": "产品名称",
  "category": "产品类别",
  "tagline": "一句话标语（不超过 20 字）",
  "description": "产品详细描述（50-200 字）",
  "features": [{{"name": "功能名", "description": "功能说明"}}],
  "tech_params": [{{"name": "参数名", "value": "参数值"}}],
  "target_customers": ["目标客户行业"],
  "pricing": "定价信息",
  "competitors": ["竞品名称"],
  "selling_points": ["核心卖点关键词"]
}}

只输出 JSON，不要其他内容。"""


def parse_product_from_docx(
    file_bytes: bytes, provider: LLMProvider | None = None
) -> dict[str, Any]:
    """解析 .docx，返回结构化产品草稿 + 元信息。

    返回字典结构：
    {
        "product": { ...ProductBase 字段... },
        "char_count": 文档字符数,
        "extractor": "llm" 或 "heuristic",
        "note": "解析说明"
    }
    """
    if provider is None:
        provider = get_provider()

    blocks = extract_blocks_from_docx(file_bytes)
    text = blocks_to_text(blocks)
    char_count = len(text)

    # 空文档直接返回空草稿
    if char_count == 0:
        return {
            "product": _empty_product(),
            "char_count": 0,
            "extractor": "none",
            "note": "文档内容为空，请手动录入。",
        }

    # 真实 LLM：让大模型抽取
    if provider.name != "mock-engine":
        product, note = _extract_with_llm(provider, text)
        return {
            "product": product,
            "char_count": char_count,
            "extractor": "llm",
            "note": note,
        }

    # Mock 模式：启发式规则抽取
    product, note = _extract_heuristic(blocks, text)
    return {
        "product": product,
        "char_count": char_count,
        "extractor": "heuristic",
        "note": note,
    }


def _extract_with_llm(provider: LLMProvider, text: str) -> tuple[dict, str]:
    """用 LLM 从手册文本中提取结构化产品信息。

    返回 (product_dict, note)
    """
    # 文本过长时截断，避免超出 LLM 上下文（保留开头，开头通常含产品名/定位）
    truncated = text[:4000]
    prompt = _EXTRACT_PROMPT.format(doc_text=truncated)
    raw = provider.chat(
        system_prompt="你是产品信息提取助手。只输出 JSON，不要其他内容。",
        user_prompt=prompt,
        temperature=0.1,    # 极低温度 → 抽取结果更稳定、不发散
    )
    try:
        obj = _parse_json(raw)
        product = _normalize_product(obj)
        return product, "已由 LLM 从手册中提取，请核对后保存。"
    except (json.JSONDecodeError, ValueError):
        # LLM 返回异常 → 回退到启发式
        blocks = _text_to_blocks(truncated)
        product, note = _extract_heuristic(blocks, truncated)
        return product, "LLM 解析失败，已回退到规则提取。" + note


def _extract_heuristic(blocks: list[dict], text: str) -> tuple[dict, str]:
    """Mock 模式：用标题层级 + 关键词规则启发式抽取。

    规则：
    - 产品名：第一个标题；没有则取第一段前 20 字
    - 描述：前 2-3 个非标题段落拼起来
    - 功能：每个标题 + 紧跟其后的段落 → 一个 feature
    - 技术参数：含"："且值简短的行，或含 版本/部署/并发 等关键词的行
    - 定价：含 ¥/元/价格/定价 的行
    - 卖点：含 优势/卖点/亮点/特点 的行
    - 竞品：含 竞品/对比 的行
    """
    product = _empty_product()
    found: list[str] = []

    # 产品名
    headings = [b for b in blocks if b["type"] == "heading"]
    paragraphs = [b for b in blocks if b["type"] == "paragraph"]
    if headings:
        product["name"] = headings[0]["text"][:50]
        found.append("名称")
    elif paragraphs:
        product["name"] = paragraphs[0]["text"][:20]
        found.append("名称")

    # 描述：前 3 个段落
    if paragraphs:
        product["description"] = " ".join(p["text"] for p in paragraphs[:3])[:300]
        found.append("描述")

    # 功能：标题 + 紧邻段落
    features = []
    for i, b in enumerate(blocks):
        if b["type"] == "heading" and i > 0:
            # 找标题后的第一个段落作为功能说明
            desc = ""
            for j in range(i + 1, len(blocks)):
                if blocks[j]["type"] == "paragraph":
                    desc = blocks[j]["text"][:120]
                    break
                if blocks[j]["type"] == "heading":
                    break
            features.append({"name": b["text"][:30], "description": desc})
    if features:
        product["features"] = features[:8]
        found.append("功能")

    # 技术参数 / 定价 / 卖点 / 竞品：按关键词在所有文本行里找
    all_lines = [b["text"] for b in blocks]
    tech_params = []
    selling_points = []
    competitors: list[str] = []
    pricing = ""

    for line in all_lines:
        # 定价
        if not pricing and re.search(r"[¥￥]|\d+元|定价|价格|订阅", line):
            pricing = line[:60]
            found.append("定价")
        # 卖点：含 优势/卖点/亮点 等关键词的行（优先于技术参数判断）
        is_selling_line = bool(re.search(r"优势|卖点|亮点|特点|核心优势", line)) and len(line) < 50
        if is_selling_line:
            # 去掉"核心优势："之类的前缀，取冒号后的内容再拆分
            content = line.split("：", 1)[1] if "：" in line else line
            content = content.split(":", 1)[1] if ":" in content else content
            for sp in re.split(r"[，,、；;]", content):
                sp = sp.strip()
                if 2 <= len(sp) <= 12 and sp not in selling_points:
                    selling_points.append(sp)
            continue   # 卖点行不再当技术参数
        # 技术参数：形如"键：值"且值不太长，且不是定价/竞品行
        if ("：" in line or ":" in line) and len(line) < 40 \
                and not re.search(r"[¥￥]|\d+元|定价|价格|订阅|竞品|对比|竞对", line):
            sep = "：" if "：" in line else ":"
            k, _, v = line.partition(sep)
            if k.strip() and v.strip():
                tech_params.append({"name": k.strip()[:20], "value": v.strip()[:40]})
        # 竞品
        if re.search(r"竞品|对比|竞对", line):
            for c in re.split(r"[，,、和与]", line):
                c = re.sub(r"竞品|对比|竞对|：|:", "", c).strip()
                if 2 <= len(c) <= 12 and c not in competitors:
                    competitors.append(c)

    if tech_params:
        product["tech_params"] = tech_params[:10]
        found.append("技术参数")
    if pricing:
        product["pricing"] = pricing
    if selling_points:
        product["selling_points"] = selling_points[:6]
        found.append("卖点")
    if competitors:
        product["competitors"] = competitors[:6]
        found.append("竞品")

    note = f"规则提取完成，识别到：{ '、'.join(found) if found else '未识别到关键字段' }。请核对补全后保存。"
    return product, note


# ===== 工具函数 =====

def _empty_product() -> dict:
    """返回空的产品草稿（字段结构与 ProductBase 对齐）。"""
    return {
        "name": "",
        "category": "",
        "tagline": "",
        "description": "",
        "features": [],
        "tech_params": [],
        "target_customers": [],
        "pricing": "",
        "competitors": [],
        "selling_points": [],
    }


def _normalize_product(obj: dict) -> dict:
    """把 LLM 返回的 JSON 规范化为产品草稿。

    主要是保证各字段类型正确（LLM 偶尔会把数组返回成字符串），
    缺失字段补默认空值。
    """
    base = _empty_product()
    base.update({k: v for k, v in obj.items() if k in base})
    # 确保列表字段确实是列表
    for key in ("features", "tech_params", "target_customers", "competitors", "selling_points"):
        if not isinstance(base.get(key), list):
            base[key] = []
    # features / tech_params 元素补全键
    base["features"] = [
        {"name": f.get("name", ""), "description": f.get("description", "")}
        for f in base["features"] if isinstance(f, dict)
    ]
    base["tech_params"] = [
        {"name": t.get("name", ""), "value": t.get("value", "")}
        for t in base["tech_params"] if isinstance(t, dict)
    ]
    return base


def _parse_json(raw: str) -> dict:
    """解析 LLM 返回的 JSON，兼容 ```json 代码块包裹。"""
    text = raw.strip()
    if text.startswith("```"):
        # 去掉首尾的 ``` 围栏
        text = text.split("```", 2)
        # text[1] 是围栏内的内容，可能以 "json" 开头
        inner = text[1] if len(text) > 1 else ""
        if inner.startswith("json"):
            inner = inner[4:]
        text = inner.strip()
    return json.loads(text)


def _text_to_blocks(text: str) -> list[dict]:
    """纯文本回退成块列表（LLM 失败时用）。

    把以 ## 开头的行当作标题，其余当作段落。
    """
    blocks = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            blocks.append({"type": "heading", "level": 2, "text": line[3:].strip()})
        else:
            blocks.append({"type": "paragraph", "text": line})
    return blocks
