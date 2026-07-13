"""生成历史管理 API：列表、详情、编辑、删除、导出。

对应需求：F4-1 ~ F4-5

每次调用 /api/generate 生成的内容都会保存到 history 表，
用户可以回看历史、编辑内容、导出为 Markdown/TXT 文件。
"""
from fastapi import APIRouter, HTTPException, Response, Depends, Query
from psycopg2.extras import Json
import base64
import io
import re
from app.database import query, query_one, transaction, _parse_json_fields
from app.models import HistoryItem, HistoryUpdate, FeedbackRequest
from app.auth import get_current_user, is_owner_or_admin
from app.seo_analyzer import analyze as seo_analyze

router = APIRouter(prefix="/api/history", tags=["生成历史"])

# history 表中的 JSONB 字段
JSON_FIELDS = ["params", "versions", "agent_trace", "issues", "feedback_voters"]


def _inject_current_feedback(rows, user_id: str):
    """把 feedback_voters 中当前用户的态度回填到 feedback 字段，并统计赞/踩总数。

    feedback 单值会被多人覆盖，改用 feedback_voters 记录每个成员的态度；
    返回前用当前用户的态度回填 feedback（前端读 row.feedback 仍是当前用户视角），
    并从 feedback_voters 统计 like_count / dislike_count 供前端按钮显示数量。
    支持单条 dict 或列表。
    """
    if isinstance(rows, dict):
        rows = [rows]
    for r in rows:
        voters = r.get("feedback_voters") or {}
        if isinstance(voters, dict) and voters:
            r["feedback"] = voters.get(user_id)
            r["like_count"] = sum(1 for v in voters.values() if v == "like")
            r["dislike_count"] = sum(1 for v in voters.values() if v == "dislike")
        else:
            r["like_count"] = 0
            r["dislike_count"] = 0
    return rows


@router.get("", response_model=list[HistoryItem])
@router.get("/", response_model=list[HistoryItem], include_in_schema=False)
def list_history(
    member: str | None = Query(None, description="按创建人(member id)筛选"),
    limit: int = Query(200, ge=1, le=500, description="返回条数上限（默认 200，避免历史积压后全表加载）"),
    offset: int = Query(0, ge=0, description="偏移量，用于分页"),
    user: dict = Depends(get_current_user),
):
    """查询历史记录列表（按时间倒序）。全员可见（共享），可按创建人筛选。

    默认返回最近 200 条，避免历史积压后全表加载拖慢响应；可用 limit/offset 翻页。
    """
    if member:
        rows = query(
            "SELECT * FROM history WHERE created_by = %s ORDER BY created_at DESC LIMIT %s OFFSET %s",
            (member, limit, offset),
        )
    else:
        rows = query(
            "SELECT * FROM history ORDER BY created_at DESC LIMIT %s OFFSET %s",
            (limit, offset),
        )
    rows = [_parse_json_fields(r, JSON_FIELDS) for r in rows]
    return _inject_current_feedback(rows, user["id"])


@router.get("/{history_id}", response_model=HistoryItem)
def get_history(history_id: str, user: dict = Depends(get_current_user)):
    """查询单条历史记录详情。"""
    row = query_one("SELECT * FROM history WHERE id = %s", (history_id,))
    if not row:
        raise HTTPException(404, f"历史记录 {history_id} 不存在")
    row = _parse_json_fields(row, JSON_FIELDS)
    _inject_current_feedback(row, user["id"])
    return row


@router.put("/{history_id}", response_model=HistoryItem)
def update_history(history_id: str, body: HistoryUpdate, user: dict = Depends(get_current_user)):
    """编辑历史记录中的内容版本。仅创建者或管理员可改。

    用户在前端手动修改生成的文案后，通过此接口保存。
    """
    existing = query_one("SELECT id, created_by FROM history WHERE id = %s", (history_id,))
    if not existing:
        raise HTTPException(404, f"历史记录 {history_id} 不存在")
    if not is_owner_or_admin(user, existing.get("created_by")):
        raise HTTPException(403, "无权编辑他人的历史记录")
    if body.versions is not None:
        # 内容可能被用户编辑过，服务端重新计算 SEO（纯规则引擎，毫秒级）
        for v in body.versions:
            v.seo = seo_analyze(v.title, v.body)
        with transaction() as cur:
            cur.execute(
                "UPDATE history SET versions = %s WHERE id = %s",
                (Json([v.model_dump() for v in body.versions]), history_id),
            )
    return get_history(history_id, user)


@router.delete("/{history_id}")
def delete_history(history_id: str, user: dict = Depends(get_current_user)):
    """删除历史记录。仅创建者或管理员可删。"""
    existing = query_one("SELECT id, created_by FROM history WHERE id = %s", (history_id,))
    if not existing:
        raise HTTPException(404, f"历史记录 {history_id} 不存在")
    if not is_owner_or_admin(user, existing.get("created_by")):
        raise HTTPException(403, "无权删除他人的历史记录")
    with transaction() as cur:
        cur.execute("DELETE FROM history WHERE id = %s", (history_id,))
    return {"message": f"历史记录 {history_id} 已删除"}


# ===== 导出辅助：图片 data URL 解析 + SVG 转 PNG + docx 构建 =====

def _parse_data_url(data_url: str):
    """解析 data URL，返回 (mime, 图片字节) 或 (None, None)。"""
    if not data_url or not data_url.startswith("data:"):
        return None, None
    m = re.match(r"data:([\w+\-.]+/[\w+\-.]+);base64,(.*)", data_url, re.S)
    if not m:
        return None, None
    try:
        return m.group(1), base64.b64decode(m.group(2))
    except Exception:
        return None, None


def _build_docx(data: dict) -> bytes:
    """把历史记录导出为 docx（含配图），返回字节。

    配图是 data URL（SVG/PNG）或 http URL；SVG 需转 PNG 才能插入 docx；
    下载/转换/插入失败则跳过该图，不影响文案导出。
    markdown 标题/列表转为 docx 对应样式，表格按行原样输出。
    配图按小节边界穿插在正文中（与前端 renderArticle 一致），而非全部堆在末尾。
    """
    from docx import Document

    doc = Document()
    doc.add_heading(f"{data['product_name']} - {data['scenario_name']}", level=0)
    doc.add_paragraph(f"渠道：{data['channel']} | 风格：{data['style']} | 生成时间：{data['created_at']}")

    for v in data.get("versions", []):
        title = v.get("title", "")
        doc.add_heading(f"版本 {v.get('index', 1)}：{title}", level=1)
        body = v.get("body") or ""
        images = v.get("images") or ([{"url": v["image"]}] if v.get("image") else [])

        # 正文按空行切成块
        blocks = [b for b in re.split(r"\n\s*\n", body) if b.strip()]
        # 去掉与标题重复的正文首标题（与前端 renderArticle 一致，如 mock 的 ## 产品名）
        if blocks and title:
            m = re.match(r"^#{1,2}\s+(.*)$", blocks[0].lstrip())
            if m:
                h = m.group(1).strip()
                if h and (h in title or title in h):
                    blocks.pop(0)

        # 计算配图穿插点（块索引，在其后插入）：首块 + 标题块 + 末块，去重保序
        points = []

        def push_point(i):
            if 0 <= i < len(blocks) and i not in points:
                points.append(i)

        if blocks:
            push_point(0)
            for i, b in enumerate(blocks):
                if b.lstrip().startswith("#"):
                    push_point(i)
            push_point(len(blocks) - 1)

        # 每张图分配一个穿插点（不足则循环复用同一位置）
        after = {}
        if points:
            for k, img in enumerate(images):
                after.setdefault(points[k % len(points)], []).append(img)
        else:
            after[-1] = list(images)    # 无正文块：配图全部放末尾

        def add_line(s):
            """把一行 markdown 转为 docx 段落/标题/列表项。"""
            if s.startswith("###"):
                doc.add_heading(s.lstrip("# ").strip(), level=3)
            elif s.startswith("##"):
                doc.add_heading(s.lstrip("# ").strip(), level=2)
            elif s.startswith("#"):
                doc.add_heading(s.lstrip("# ").strip(), level=1)
            elif s.startswith("- ") or s.startswith("• "):
                doc.add_paragraph(s.lstrip("-• ").strip(), style="List Bullet")
            else:
                doc.add_paragraph(s)

        # 逐块渲染正文，在穿插点后插入对应配图
        for i, blk in enumerate(blocks):
            for line in blk.split("\n"):
                s = line.strip()
                if s:
                    add_line(s)
            for img in after.get(i, []):
                _insert_docx_image(
                    doc,
                    img.get("url") if isinstance(img, dict) else None,
                    img.get("caption", "") if isinstance(img, dict) else "",
                )
        # 无正文块的剩余配图
        for img in after.get(-1, []):
            _insert_docx_image(
                doc,
                img.get("url") if isinstance(img, dict) else None,
                img.get("caption", "") if isinstance(img, dict) else "",
            )
        if v.get("tags"):
            doc.add_paragraph(f"标签：{', '.join(v['tags'])}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _insert_docx_image(doc, url: str, caption: str = "") -> None:
    """把一张配图插入 docx：data URL 直接解析，http(s) URL 下载后插入；SVG 跳过；失败静默跳过。"""
    img_bytes = None
    mime = None
    if url.startswith("data:"):
        mime, img_bytes = _parse_data_url(url)
    elif url.startswith("http://") or url.startswith("https://"):
        try:
            import httpx
            r = httpx.get(url, timeout=30, follow_redirects=True)
            r.raise_for_status()
            img_bytes = r.content
            mime = r.headers.get("content-type", "")
        except Exception:
            return    # 下载失败：跳过该图，不影响整体导出
    if not img_bytes:
        return
    from docx.shared import Inches
    # python-docx 不支持 SVG（曾用 svglib 转换，但依赖未安装且配图已改 PNG 生成，这里直接跳过）
    if mime and "svg" in mime:
        return
    try:
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
        if caption:
            cap_p = doc.add_paragraph(caption)
            if cap_p.runs:
                cap_p.runs[0].italic = True
        doc.add_paragraph()
    except Exception:
        pass


@router.get("/{history_id}/export")
def export_history(history_id: str, format: str = "markdown", user: dict = Depends(get_current_user)):
    """导出历史记录为文件。

    GET /api/history/H001/export?format=markdown → 下载 .md 文件
    GET /api/history/H001/export?format=txt      → 下载 .txt 文件

    中文文件名需要特殊处理（Content-Disposition 编码），
    否则浏览器下载时中文会乱码。
    """
    row = query_one("SELECT * FROM history WHERE id = %s", (history_id,))
    if not row:
        raise HTTPException(404, f"历史记录 {history_id} 不存在")
    data = _parse_json_fields(row, JSON_FIELDS)
    versions = data.get("versions", [])

    # docx 导出（含配图，二进制，单独返回）
    if format == "docx":
        content_bytes = _build_docx(data)
        from urllib.parse import quote
        safe_name = data["product_name"].replace(" ", "_")
        filename = f"{safe_name}_{history_id}.docx"
        ascii_fallback = f"shenxing_{history_id}.docx"
        return Response(
            content=content_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": (
                    f"attachment; filename={ascii_fallback}; "
                    f"filename*=UTF-8''{quote(filename)}"
                )
            },
        )

    if format == "txt":
        # 纯文本格式
        lines = []
        for v in versions:
            lines.append(f"=== 版本 {v.get('index', 1)}：{v.get('title', '')} ===\n")
            lines.append(v.get("body", ""))
            lines.append("\n")
        content = "\n".join(lines)
        media_type = "text/plain"
        ext = "txt"
    else:
        # Markdown 格式
        parts = [f"# {data['product_name']} - {data['scenario_name']}\n"]
        parts.append(f"> 渠道：{data['channel']} | 风格：{data['style']} | 生成时间：{data['created_at']}\n")
        for v in versions:
            parts.append(f"\n## 版本 {v.get('index', 1)}：{v.get('title', '')}\n")
            parts.append(v.get("body", ""))
            if v.get("tags"):
                parts.append(f"\n\n**标签：** {', '.join(v['tags'])}")
            parts.append("\n")
        content = "\n".join(parts)
        media_type = "text/markdown"
        ext = "md"

    # 中文文件名编码处理
    # RFC 5987 规定：filename* 用于非 ASCII 文件名，格式为 UTF-8''编码后的文件名
    from urllib.parse import quote
    safe_name = data["product_name"].replace(" ", "_")
    filename = f"{safe_name}_{history_id}.{ext}"
    ascii_fallback = f"shenxing_{history_id}.{ext}"  # ASCII 兜底文件名
    return Response(
        content=content.encode("utf-8"),
        media_type=media_type,
        headers={
            "Content-Disposition": (
                f"attachment; filename={ascii_fallback}; "
                f"filename*=UTF-8''{quote(filename)}"
            )
        },
    )


# ===== 加分项：用户反馈（赞 / 踩）=====

@router.put("/{history_id}/feedback", response_model=HistoryItem)
def set_feedback(history_id: str, body: FeedbackRequest, user: dict = Depends(get_current_user)):
    """标记用户对某条生成内容的反馈（点赞 / 踩 / 取消）。任何登录用户可操作。

    PUT /api/history/H001/feedback
    {"feedback": "like"}      → 点赞
    {"feedback": "dislike"}   → 踩
    {"feedback": ""}          → 取消标记
    """
    existing = query_one("SELECT id, feedback_voters FROM history WHERE id = %s", (history_id,))
    if not existing:
        raise HTTPException(404, f"历史记录 {history_id} 不存在")

    # 规范化取值：去空格；空字符串表示取消
    fb = (body.feedback or "").strip()
    if fb and fb not in ("like", "dislike"):
        raise HTTPException(400, "feedback 取值需为 like / dislike / 空字符串")
    # 空字符串 → 写入 NULL（数据库里用 NULL 表示"未标记"）
    # per-user 改票/取消：更新 feedback_voters 中当前用户的态度（互不覆盖）
    voters = existing.get("feedback_voters") or {}
    if not isinstance(voters, dict):
        voters = {}
    if fb:
        voters[user["id"]] = fb
    else:
        voters.pop(user["id"], None)

    with transaction() as cur:
        cur.execute(
            "UPDATE history SET feedback_voters = %s WHERE id = %s",
            (Json(voters), history_id),
        )
    return get_history(history_id, user)